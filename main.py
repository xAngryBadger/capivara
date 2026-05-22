from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import io
import os

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter

FONT_DIR = "/usr/share/fonts/TTF"
for name, file in [
    ("DejaVu", "DejaVuSans.ttf"),
    ("DejaVu-Bold", "DejaVuSans-Bold.ttf"),
    ("DejaVuMono", "DejaVuSansMono.ttf"),
]:
    path = os.path.join(FONT_DIR, file)
    if os.path.exists(path):
        pdfmetrics.registerFont(TTFont(name, path))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _pdf_stream(writer: PdfWriter) -> StreamingResponse:
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf")


def _get_styles():
    base = getSampleStyleSheet()
    font = "DejaVu" if "DejaVu" in pdfmetrics._fonts else "Helvetica"
    bold = "DejaVu-Bold" if "DejaVu-Bold" in pdfmetrics._fonts else "Helvetica-Bold"
    mono = "DejaVuMono" if "DejaVuMono" in pdfmetrics._fonts else "Courier"
    return {
        "Normal": ParagraphStyle("DejaVuNormal", parent=base["Normal"], fontName=font, fontSize=10, leading=14),
        "Heading1": ParagraphStyle("DejaVuH1", parent=base["Heading1"], fontName=bold, fontSize=14, leading=18, spaceAfter=6),
        "Code": ParagraphStyle("DejaVuCode", parent=base["Code"], fontName=mono, fontSize=8, leading=11),
    }


async def _read_upload(file: UploadFile) -> bytes:
    return await file.read()


# --- Converters ---


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    buf_in = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()
    doc = Document(buf_in)
    doc_pdf = SimpleDocTemplate(buf_out, pagesize=A4)
    styles = _get_styles()
    story = []
    for para in doc.paragraphs:
        story.append(Paragraph(para.text, styles["Normal"]))
        story.append(Spacer(1, 6))
    doc_pdf.build(story)
    buf_out.seek(0)
    return buf_out.getvalue()


def convert_xlsx_to_pdf(xlsx_bytes: bytes) -> bytes:
    buf_in = io.BytesIO(xlsx_bytes)
    buf_out = io.BytesIO()
    wb = load_workbook(buf_in, data_only=True)
    doc_pdf = SimpleDocTemplate(buf_out, pagesize=A4)
    styles = _get_styles()
    story = []
    for sheet in wb.worksheets:
        story.append(Paragraph(f"Sheet: {sheet.title}", styles["Heading1"]))
        story.append(Spacer(1, 6))
        max_row = sheet.max_row or 0
        max_col = sheet.max_column or 0
        for row in sheet.iter_rows(min_row=1, max_row=min(50, max_row), min_col=1, max_col=min(20, max_col)):
            row_text = " | ".join(str(cell.value) if cell.value else "" for cell in row)
            story.append(Paragraph(row_text, styles["Code"]))
            story.append(Spacer(1, 3))
        story.append(Spacer(1, 12))
    doc_pdf.build(story)
    buf_out.seek(0)
    return buf_out.getvalue()


# --- PDF operations (pypdf only) ---


def merge_pdfs(pdf_bytes_list: list[bytes]) -> bytes:
    writer = PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def split_pdf(pdf_bytes: bytes, pages: str) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    total = len(reader.pages)

    selected = set()
    for part in pages.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            for p in range(int(start), int(end) + 1):
                if 1 <= p <= total:
                    selected.add(p - 1)
        else:
            p = int(part)
            if 1 <= p <= total:
                selected.add(p - 1)

    for idx in sorted(selected):
        writer.add_page(reader.pages[idx])

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def rotate_pdf(pdf_bytes: bytes, angle: int) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def compress_pdf(pdf_bytes: bytes) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def protect_pdf(pdf_bytes: bytes, password: str, restrict_print: bool = False) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def unlock_pdf(pdf_bytes: bytes, password: str = "") -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        reader.decrypt(password)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


# --- Overlay operations (reportlab + pypdf) ---


def _overlay_on_pdf(pdf_bytes: bytes, overlay_func) -> bytes:
    from reportlab.pdfgen import canvas

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    page_count = len(reader.pages)

    for i, page in enumerate(reader.pages):
        overlay_buf = io.BytesIO()
        page_w = float(page.mediabox.width)
        page_h = float(page.mediabox.height)
        c = canvas.Canvas(overlay_buf, pagesize=(page_w, page_h))
        overlay_func(c, page_w, page_h, i + 1, page_count)
        c.save()
        overlay_buf.seek(0)

        overlay_reader = PdfReader(overlay_buf)
        page.merge_page(overlay_reader.pages[0])
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def watermark_pdf(pdf_bytes: bytes, text: str, opacity: float = 0.3, font_size: int = 48) -> bytes:
    import math

    def draw(c, w, h, _page_num, _total):
        c.saveState()
        c.setFillAlpha(opacity)
        c.setFillGray(0.5)
        c.setFont("DejaVu" if "DejaVu" in pdfmetrics._fonts else "Helvetica", font_size)
        c.translate(w / 2, h / 2)
        c.rotate(45)
        c.drawCentredString(0, 0, text)
        c.restoreState()

    return _overlay_on_pdf(pdf_bytes, draw)


def page_numbers_pdf(pdf_bytes: bytes, position: str = "bottom-center", start_at: int = 1) -> bytes:
    positions = {
        "bottom-center": lambda w, h: (w / 2, 30),
        "bottom-right": lambda w, h: (w - 60, 30),
        "top-center": lambda w, h: (w / 2, h - 30),
        "top-right": lambda w, h: (w - 60, h - 30),
    }
    pos_fn = positions.get(position, positions["bottom-center"])

    def draw(c, w, h, page_num, _total):
        x, y = pos_fn(w, h)
        c.setFont("DejaVu" if "DejaVu" in pdfmetrics._fonts else "Helvetica", 10)
        c.setFillGray(0.4)
        c.drawCentredString(x, y, str(page_num + start_at - 1))

    return _overlay_on_pdf(pdf_bytes, draw)


def header_footer_pdf(pdf_bytes: bytes, header: str = "", footer: str = "") -> bytes:
    def draw(c, w, h, _page_num, _total):
        c.setFont("DejaVu" if "DejaVu" in pdfmetrics._fonts else "Helvetica", 9)
        c.setFillGray(0.4)
        if header:
            c.drawCentredString(w / 2, h - 30, header)
        if footer:
            c.drawCentredString(w / 2, 30, footer)

    return _overlay_on_pdf(pdf_bytes, draw)


# --- PyMuPDF operations (lazy) ---


def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    from docx import Document as DocxDocument
    docx_doc = DocxDocument()
    for para_text in text.split("\n"):
        if para_text.strip():
            docx_doc.add_paragraph(para_text)

    buf = io.BytesIO()
    docx_doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# --- OCR (lazy, requires tesseract + poppler) ---


def ocr_pdf(pdf_bytes: bytes, lang: str = "por") -> bytes:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image

    images = convert_from_bytes(pdf_bytes, dpi=200)
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()

    for i, img in enumerate(images):
        ocr_text = pytesseract.image_to_string(img, lang=lang)

        overlay_buf = io.BytesIO()
        from reportlab.pdfgen import canvas
        if i < len(reader.pages):
            page = reader.pages[i]
        else:
            continue
        page_w = float(page.mediabox.width)
        page_h = float(page.mediabox.height)
        c = canvas.Canvas(overlay_buf, pagesize=(page_w, page_h))
        c.setFillAlpha(0)
        c.setFont("DejaVu" if "DejaVu" in pdfmetrics._fonts else "Helvetica", 1)
        for j, line in enumerate(ocr_text.split("\n")):
            y = page_h - 20 - j * 14
            if y < 20:
                break
            c.drawString(20, y, line)
        c.save()
        overlay_buf.seek(0)
        overlay_reader = PdfReader(overlay_buf)
        page.merge_page(overlay_reader.pages[0])
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


# --- PDF to Images (lazy, requires PyMuPDF) ---


def pdf_to_images(pdf_bytes: bytes, fmt: str = "png", dpi: int = 200) -> bytes:
    import fitz
    import zipfile

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as zf:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            ext = "png" if fmt == "png" else "jpeg"
            img_bytes = pix.tobytes(ext)
            zf.writestr(f"page_{i + 1:03d}.{ext}", img_bytes)
    doc.close()

    zip_buf.seek(0)
    return zip_buf.getvalue()


# --- PDF/A (lazy, requires pikepdf) ---


def pdfa_convert(pdf_bytes: bytes, level: str = "PDF/A-1b") -> bytes:
    import pikepdf

    pdf = pikepdf.open(stream=io.BytesIO(pdf_bytes))
    with pdf.open_metadata() as meta:
        meta["xmp:CreatorTool"] = "Capivara PDF Suite"
    buf = io.BytesIO()
    pdf.save(buf)
    pdf.close()
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# ENDPOINTS
# ============================================================


@app.get("/")
async def root():
    return {"message": "Capivara PDF Suite API", "tools": 15}


@app.post("/api/convert/docx-pdf")
async def api_docx_pdf(file: UploadFile = File(...)):
    try:
        data = await _read_upload(file)
        result = convert_docx_to_pdf(data)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert/xlsx-pdf")
async def api_xlsx_pdf(file: UploadFile = File(...)):
    try:
        data = await _read_upload(file)
        result = convert_xlsx_to_pdf(data)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert/pdf-compress")
async def api_compress(file: UploadFile = File(...)):
    try:
        data = await _read_upload(file)
        result = compress_pdf(data)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/merge")
async def api_merge(files: list[UploadFile] = File(...)):
    try:
        pdf_bytes_list = []
        for f in files:
            pdf_bytes_list.append(await _read_upload(f))
        result = merge_pdfs(pdf_bytes_list)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/split")
async def api_split(file: UploadFile = File(...), pages: str = Form("1")):
    try:
        data = await _read_upload(file)
        result = split_pdf(data, pages)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/rotate")
async def api_rotate(file: UploadFile = File(...), angle: str = Form("90")):
    try:
        data = await _read_upload(file)
        result = rotate_pdf(data, int(angle))
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/watermark")
async def api_watermark(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENCIAL"),
    opacity: str = Form("0.3"),
    font_size: str = Form("48"),
):
    try:
        data = await _read_upload(file)
        result = watermark_pdf(data, text, float(opacity), int(font_size))
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/page-numbers")
async def api_page_numbers(
    file: UploadFile = File(...),
    position: str = Form("bottom-center"),
    start_at: str = Form("1"),
):
    try:
        data = await _read_upload(file)
        result = page_numbers_pdf(data, position, int(start_at))
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/header-footer")
async def api_header_footer(
    file: UploadFile = File(...),
    header: str = Form(""),
    footer: str = Form(""),
):
    try:
        data = await _read_upload(file)
        result = header_footer_pdf(data, header, footer)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/protect")
async def api_protect(
    file: UploadFile = File(...),
    password: str = Form(...),
    restrict_print: str = Form(""),
):
    try:
        data = await _read_upload(file)
        result = protect_pdf(data, password, bool(restrict_print))
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/unlock")
async def api_unlock(
    file: UploadFile = File(...),
    password: str = Form(""),
):
    try:
        data = await _read_upload(file)
        result = unlock_pdf(data, password)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert/pdf-docx")
async def api_pdf_docx(file: UploadFile = File(...)):
    try:
        data = await _read_upload(file)
        result = pdf_to_docx(data)
        return StreamingResponse(
            io.BytesIO(result),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/ocr")
async def api_ocr(file: UploadFile = File(...), lang: str = Form("por")):
    try:
        data = await _read_upload(file)
        result = ocr_pdf(data, lang)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/pdf-to-images")
async def api_pdf_to_images(
    file: UploadFile = File(...),
    format: str = Form("png"),
    dpi: str = Form("200"),
):
    try:
        data = await _read_upload(file)
        result = pdf_to_images(data, format, int(dpi))
        return StreamingResponse(io.BytesIO(result), media_type="application/zip")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/pdfa")
async def api_pdfa(file: UploadFile = File(...), level: str = Form("PDF/A-1b")):
    try:
        data = await _read_upload(file)
        result = pdfa_convert(data, level)
        return StreamingResponse(io.BytesIO(result), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert/batch")
async def api_batch(files: list[UploadFile] = File(...)):
    try:
        results = []
        for f in files:
            try:
                data = await _read_upload(f)
                if f.filename and f.filename.endswith(".docx"):
                    convert_docx_to_pdf(data)
                    results.append({"original": f.filename, "status": "success"})
                elif f.filename and f.filename.endswith(".xlsx"):
                    convert_xlsx_to_pdf(data)
                    results.append({"original": f.filename, "status": "success"})
                else:
                    results.append({"original": f.filename, "status": "error", "error": "Formato não suportado"})
            except Exception as e:
                results.append({"original": f.filename, "status": "error", "error": str(e)})
        return {"results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8001)))
